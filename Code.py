import streamlit as st
import datetime
from docx import Document
from io import BytesIO
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

st.title("Travel Itinerary Generator")

# Upload image
uploaded_image = "logo.jpg"  # Set to path or allow user to upload an image

# --- Passenger Details ---
st.header("Passenger Details")
col1, col2 = st.columns(2)
with col1:
    lead_passenger = st.text_input("Lead Passenger", "Mr. Stephen Xavier")
with col2:
    number_of_pax = st.number_input("Number of Passengers", min_value=1, value=2)
tfg_reference = st.text_input("TFG Reference", "TFG Reference")

# --- Flight Schedule ---
st.header("Flight Schedule")
arrival_flight = st.text_input("Arrival Flight", "EK 654 at 1555 Hrs on 29 Jan 2026")
departure_flight = st.text_input("Departure Flight", "EK 655 at 2205 Hrs on 15 Feb 2026")
col1, col2 = st.columns(2)
with col1:
    travel_date = st.date_input("Travel Date", datetime.date(2026, 1, 29))
with col2:
    departure_date = st.date_input("Departure Date", datetime.date(2026, 2, 15))
number_of_days = st.text_input("Number of Days", "Sri Lanka – 17 Nights / 18 Days")

# --- Accommodation Schedule ---
st.header("Accommodation Schedule")
accommodation_schedule = []
accommodation_count = st.number_input("Number of Accommodations", min_value=1, value=1)

for i in range(accommodation_count):
    st.subheader(f"Accommodation {i + 1}")
    col1, col2 = st.columns(2)
    with col1:
        check_in_date = st.date_input(f"Check-in Date {i + 1}", datetime.date(2026, 1, 29))
    with col2:
        check_out_date = st.date_input(f"Check-out Date {i + 1}", datetime.date(2026, 1, 30))
    location = st.text_input(f"Location {i + 1}", "The Wallawwa, Kotugoda")
    details = st.text_input(f"Details {i + 1}", "1 night – 1 DBL Garden Suite on BB basis")
    notes = st.text_area(f"Notes {i + 1}", "10% EBO applied if booked 90 days prior to check-in")
    
    accommodation_schedule.append({
        "check_in_date": check_in_date,
        "check_out_date": check_out_date,
        "location": location,
        "details": details,
        "notes": notes
    })

# --- Activity and Transfer Schedule ---
st.header("Activity and Transfer Schedule")
activity_schedule = []
activity_count = st.number_input("Number of Activities/Transfers", min_value=1, value=1)

for i in range(activity_count):
    st.subheader(f"Activity/Transfer {i + 1}")
    col1, col2 = st.columns(2)
    with col1:
        date = st.date_input(f"Date {i + 1}", datetime.date(2026, 1, 29))
    with col2:
        time = st.text_input(f"Time {i + 1}", "Silk Route Arrival")
    description = st.text_area(f"Description {i + 1}", "Transfer to Kotugoda (Approx 30 Minutes)")
    
    activity_schedule.append({
        "date": date,
        "time": time,
        "description": description
    })

# --- Generate Word Document ---
def generate_word_doc():
    doc = Document()
    
    # Add image to the top-right
    if uploaded_image:
        image = Image.open(uploaded_image)
        image_stream = BytesIO()
        image.save(image_stream, format="PNG")
        image_stream.seek(0)
        
        # Add image to header
        section = doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(3))  # Adjust width as needed
        
        # Align image to the right
        paragraph.alignment = 2  # Right align

    doc.add_heading("Travel Itinerary", level=1)
    
    doc.add_heading("Passenger Details", level=2)
    doc.add_paragraph(f"Lead Passenger: {lead_passenger}")
    doc.add_paragraph(f"Number of Passengers: {number_of_pax}")
    doc.add_paragraph(f"TFG Reference: {tfg_reference}")
    
    doc.add_heading("Flight Schedule", level=2)
    doc.add_paragraph(f"Arrival Flight: {arrival_flight}")
    doc.add_paragraph(f"Departure Flight: {departure_flight}")
    doc.add_paragraph(f"Travel Date: {travel_date.strftime('%d %B %Y')}")
    doc.add_paragraph(f"Departure Date: {departure_date.strftime('%d %B %Y')}")
    doc.add_paragraph(f"Number of Days: {number_of_days}")
    
    doc.add_heading("Accommodation Schedule", level=2)
    for acc in accommodation_schedule:
        doc.add_paragraph(f"{acc['check_in_date'].strftime('%d %B %Y')} - {acc['check_out_date'].strftime('%d %B %Y')}: {acc['location']}")
        doc.add_paragraph(f"  Details: {acc['details']}")
        if acc['notes']:
            doc.add_paragraph(f"  Notes: {acc['notes']}")
    
    doc.add_heading("Activity and Transfer Schedule", level=2)
    for act in activity_schedule:
        doc.add_paragraph(f"{act['date'].strftime('%d %B %Y')} - {act['time']}: {act['description']}")
    
    # Save document to in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Download Button ---
st.header("Download Itinerary")
if st.button("Generate Itinerary"):
    word_file = generate_word_doc()
    st.download_button(label="Click here to download",
                       data=word_file,
                       file_name="Travel_Itinerary.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
